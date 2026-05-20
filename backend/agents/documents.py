"""
Document ingestion + risk-highlighting for NXT8 Compliance persona.

Pipeline:
1.  Receive PDF / DOCX / TXT upload (≤ 10 MB).
2.  Extract text (pypdf / python-docx / utf-8 decode).
3.  Chunk to ≤ 6 KB blocks, store each in MemPalace wing="documents"
    room=<document_id>. This makes the doc body searchable across the
    whole platform (search_memory / mempalace_search will surface it).
4.  Run a single DeepSeek pass with a "compliance reviewer" prompt and
    parse the JSON risk verdict (severity, findings, recommended_actions).
5.  Persist a `documents` collection row with the parsed result + audit
    fields so the UI can list previously uploaded docs.

Pure backend module — exposed via /api/documents/* endpoints in server.py.
"""
from __future__ import annotations

import io
import json
import logging
import re
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from agents import mempalace_bridge as mempalace_agent
from core.db import get_db
from core.deepseek import get_deepseek

logger = logging.getLogger("nxt8.documents")

MAX_BYTES = 10 * 1024 * 1024  # 10 MB
CHUNK_CHARS = 6000


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


# =====================================================================
# Extraction
# =====================================================================


def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(f"pypdf not available: {e}")
    reader = PdfReader(io.BytesIO(raw))
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            logger.warning("pdf page extract failed: %s", e)
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(f"python-docx not available: {e}")
    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also extract table rows
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_text(raw: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_text(filename: str, raw: bytes) -> str:
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(raw)
    if lower.endswith(".docx"):
        return _extract_docx(raw)
    if lower.endswith((".txt", ".md", ".csv", ".log")):
        return _extract_text(raw)
    # Best-effort fallback
    return _extract_text(raw)


def _chunk(text: str, max_chars: int = CHUNK_CHARS) -> List[str]:
    if not text:
        return []
    text = text.strip()
    if len(text) <= max_chars:
        return [text]
    chunks: List[str] = []
    i = 0
    while i < len(text):
        chunks.append(text[i:i + max_chars])
        i += max_chars
    return chunks


# =====================================================================
# Risk analysis via DeepSeek
# =====================================================================


_RISK_SYSTEM = (
    "Ты — Compliance Reviewer NXT8. Анализируешь корпоративный документ "
    "(договор, политика, NDA, оферта, regulatory) и подсвечиваешь риски.\n\n"
    "Жёсткий формат ответа — ОДИН JSON-объект в fenced-блоке, ничего больше:\n"
    "```json\n"
    "{\n"
    '  "severity": "low|medium|high|critical",\n'
    '  "summary": "2-4 предложения по сути документа",\n'
    '  "findings": [\n'
    '    {"category": "limitation|liability|payment|termination|data|regulatory|other",\n'
    '     "severity": "low|medium|high|critical",\n'
    '     "quote": "точная цитата фразы из документа (≤ 200 chars)",\n'
    '     "risk": "почему это риск (1-2 предложения)",\n'
    '     "recommendation": "что конкретно исправить или согласовать"}\n'
    "  ],\n"
    '  "recommended_actions": ["1-3 коротких action items"]\n'
    "}\n"
    "```\n"
    "Если документ выглядит безобидно — верни severity=low и findings=[].\n"
    "Не выдумывай. Цитата `quote` должна быть из текста как есть."
)


_JSON_RE = re.compile(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", re.IGNORECASE)


def _parse_risk_verdict(content: str) -> Dict[str, Any]:
    if not content:
        return {}
    m = _JSON_RE.search(content)
    raw = m.group(1) if m else content.strip()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        # try the largest balanced { … } slice
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            try:
                data = json.loads(raw[start:end + 1])
            except json.JSONDecodeError:
                return {"raw": content}
        else:
            return {"raw": content}
    if not isinstance(data, dict):
        return {"raw": content}
    return data


async def analyse_risk(text: str) -> Dict[str, Any]:
    if not text or not text.strip():
        return {"severity": "low", "summary": "Документ пустой.", "findings": [],
                "recommended_actions": [], "tokens_total": 0, "mock": False}
    snippet = text[:18000]  # cap context to be safe
    deepseek = get_deepseek()
    resp = await deepseek.chat(
        messages=[
            {"role": "system", "content": _RISK_SYSTEM},
            {"role": "user", "content": f"Документ для анализа:\n\n{snippet}"},
        ],
        temperature=0.2,
        max_tokens=2048,
    )
    verdict = _parse_risk_verdict(resp.get("content") or "")
    verdict.setdefault("severity", "low")
    verdict.setdefault("findings", [])
    verdict.setdefault("recommended_actions", [])
    verdict["tokens_total"] = int(resp.get("tokens_total") or 0)
    verdict["mock"] = bool(resp.get("mock"))
    verdict["provider"] = resp.get("provider")
    verdict["confidence"] = float(resp.get("confidence") or 0.7)
    return verdict


# =====================================================================
# Public entry: upload + analyse + persist
# =====================================================================


async def ingest_document(
    *,
    filename: str,
    content: bytes,
    company_id: str = "default",
    user_id: str = "anonymous",
    title: Optional[str] = None,
    notes: Optional[str] = None,
) -> Dict[str, Any]:
    if not content:
        raise ValueError("empty file")
    if len(content) > MAX_BYTES:
        raise ValueError(f"file exceeds {MAX_BYTES // (1024 * 1024)} MB limit")

    text = extract_text(filename, content)
    if not text.strip():
        raise ValueError("could not extract any text from file")

    document_id = str(uuid.uuid4())
    chunks = _chunk(text)

    # 1. store chunks in MemPalace under wing=documents/room=<doc_id>
    mp_stored = 0
    try:
        bridge = mempalace_agent.get_mempalace()
        for idx, chunk in enumerate(chunks):
            try:
                await bridge.store(
                    content=chunk,
                    wing="documents",
                    room=document_id,
                    metadata={
                        "company_id": company_id,
                        "filename": filename,
                        "title": title or filename,
                        "chunk_index": idx,
                        "total_chunks": len(chunks),
                    },
                    source="documents_upload",
                )
                mp_stored += 1
            except Exception as e:  # noqa: BLE001
                logger.warning("mempalace store chunk %s failed: %s", idx, e)
    except Exception as e:  # noqa: BLE001
        logger.warning("mempalace bridge unavailable: %s", e)

    # 2. risk analysis (single LLM pass)
    try:
        verdict = await analyse_risk(text)
    except Exception as e:  # noqa: BLE001
        logger.exception("risk analysis failed: %s", e)
        verdict = {"severity": "unknown", "error": str(e),
                   "findings": [], "recommended_actions": [],
                   "tokens_total": 0, "mock": False, "provider": None,
                   "confidence": 0.4}

    # 3. persist audit row
    doc = {
        "id": document_id,
        "company_id": company_id,
        "user_id": user_id,
        "filename": filename,
        "title": title or filename,
        "notes": notes,
        "size_bytes": len(content),
        "chunks": len(chunks),
        "mempalace_stored_chunks": mp_stored,
        "text_chars": len(text),
        "preview": text[:500],
        "severity": verdict.get("severity"),
        "summary": verdict.get("summary"),
        "findings": verdict.get("findings", []),
        "recommended_actions": verdict.get("recommended_actions", []),
        "mock": verdict.get("mock", False),
        "confidence": verdict.get("confidence", 0.7),
        "tokens_total": verdict.get("tokens_total", 0),
        "provider": verdict.get("provider"),
        "created_at": _now(),
    }
    await get_db().documents.insert_one(doc)

    # Return without _id (excluded on read)
    return {k: v for k, v in doc.items() if k != "_id"}


async def list_documents(company_id: Optional[str] = None,
                         limit: int = 50) -> List[Dict[str, Any]]:
    q: Dict[str, Any] = {}
    if company_id:
        q["company_id"] = company_id
    cursor = get_db().documents.find(q, {"_id": 0}).sort("created_at", -1).limit(limit)
    return await cursor.to_list(length=limit)


async def get_document(document_id: str) -> Optional[Dict[str, Any]]:
    return await get_db().documents.find_one({"id": document_id}, {"_id": 0})
